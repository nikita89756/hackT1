"""
Модуль парсера для обработки различных типов документов (PDF, веб-страницы, текстовые файлы).
"""

import os
import re
import requests
import PyPDF2
import urllib.request
from bs4 import BeautifulSoup
from urllib.parse import urlparse
from typing import Dict, List, Optional, Union
from pydantic import BaseModel
import io
from docx import Document

class DocumentHandle:
    """Класс для обработки различных типов документов."""
    
    def __init__(self, source: str):
        """
        Инициализация обработчика документов.
        Аргументы:
            source: Путь к файлу или URL
        """
        self.source = source
        self.content: Optional[str] = None
        
    def is_url(self) -> bool:
        """Проверяет, является ли источник URL."""
        try:
            result = urlparse(self.source)
            return all([result.scheme, result.netloc])
        except Exception:
            return False
            
    def is_pdf(self) -> bool:
        """Проверяет, является ли источник PDF файлом."""
        return self.source.lower().endswith('.pdf')

    def is_docx(self) -> bool:
        """Проверяет, является ли источник DOCX файлом."""
        return self.source.lower().endswith('.docx')
            
    def _convert_pdf_to_text(self, pdf_content: Union[str, bytes]) -> str:
        """
        Конвертирует PDF в текст.
        Аргументы:
            pdf_content: Содержимое PDF файла или путь к файлу
        Возвращает:
            str: Извлеченный текст из PDF
        """
        try:
            text_parts = []
            if isinstance(pdf_content, str):  # если это путь к файлу
                with open(pdf_content, mode='rb') as file:
                    content = file.read()
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                    for page in pdf_reader.pages:
                        text_parts.append(page.extract_text())
            else:  # если это bytes
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
                for page in pdf_reader.pages:
                    text_parts.append(page.extract_text())
            return "\n".join(text_parts).strip()
        except Exception as e:
            print(f"Ошибка конвертации PDF: {str(e)}")
            return ""

    def _convert_docx_to_text(self, docx_content: Union[str, bytes]) -> str:
        """
        Конвертирует DOCX в текст.
        Аргументы:
            docx_content: Содержимое DOCX файла или путь к файлу
        Возвращает:
            str: Извлеченный текст из DOCX
        """
        try:
            if isinstance(docx_content, str):  # если это путь к файлу
                doc = Document(docx_content)
            else:  # если это bytes
                doc = Document(io.BytesIO(docx_content))
            
            # Извлекаем текст из параграфов
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts).strip()
        except Exception as e:
            print(f"Ошибка конвертации DOCX: {str(e)}")
            return ""

    def get_content(self) -> str:
        """
        Получает содержимое документа.
        Возвращает:
            str: Извлеченный текст из документа
        """
        if self.content is not None:
            return self.content
            
        if self.is_url():
            if self.is_pdf():
                try:
                    response = requests.get(self.source, timeout=10)
                    if response.status_code == 200:
                        content = response.content
                        self.content = self._convert_pdf_to_text(content)
                    else:
                        print(f"Ошибка получения {self.source}: {response.status_code}")
                        self.content = ""
                except Exception as e:
                    print(f"Ошибка получения {self.source}: {str(e)}")
                    self.content = ""
            elif self.is_docx():
                try:
                    response = requests.get(self.source, timeout=10)
                    if response.status_code == 200:
                        content = response.content
                        self.content = self._convert_docx_to_text(content)
                    else:
                        print(f"Ошибка получения {self.source}: {response.status_code}")
                        self.content = ""
                except Exception as e:
                    print(f"Ошибка получения {self.source}: {str(e)}")
                    self.content = ""
            else:
                self.content = scrape_website(self.source)
        else:
            if self.is_pdf():
                self.content = self._convert_pdf_to_text(self.source)
            elif self.is_docx():
                self.content = self._convert_docx_to_text(self.source)
            else:
                try:
                    with open(self.source, 'r', encoding='utf-8') as file:
                        self.content = file.read()
                except Exception as e:
                    print(f"Ошибка чтения файла: {str(e)}")
                    self.content = ""
                    
        return self.content

    def save_as_txt(self, output_path: str) -> str:
        """
        Сохраняет содержимое в текстовый файл.
        Аргументы:
            output_path: Путь для сохранения файла
        Возвращает:
            str: Путь к сохраненному файлу
        """
        if not self.content:
            self.get_content()
            
        try:
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(self.content)
            return output_path
        except Exception as e:
            print(f"Ошибка сохранения файла: {str(e)}")
            return ""

def scrape_website(url: str) -> str:
    """
    Скрапинг информации с веб-страницы.
    Аргументы:
        url: URL сайта
    Возвращает:
        str: Извлеченный текст или пустая строка в случае ошибки
    """
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            html = response.text
            soup = BeautifulSoup(html, 'html.parser')
            
            # Удаляем ненужные теги
            for script in soup(["script", "style"]):
                script.decompose()
                
            text = soup.get_text()
            
            # Обработка текста
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            return ' '.join(chunk for chunk in chunks if chunk)
        else:
            print(f"Ошибка получения страницы: {response.status_code}")
            return ""
    except Exception as e:
        print(f"Ошибка при скрапинге {url}: {str(e)}")
        return ""
